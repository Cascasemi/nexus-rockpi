import time
import os
import subprocess
import threading
import speech_recognition as sr
from queue import Queue, Empty
from vosk import Model, KaldiRecognizer, SetLogLevel
from google import genai
from google.genai import types, errors as genai_errors
from docx import Document
from datetime import datetime
import requests
import json
from urllib.parse import urlparse
from dotenv import load_dotenv

try:
    import gpiod
except ImportError:
    gpiod = None  # libgpiod python bindings not installed — button falls back to keyboard trigger

load_dotenv()

# ============================================================
# CONFIG
# ============================================================
BUTTON_GPIO_CHIP = "/dev/gpiochip0"   # Rock Pi 4C+ typically exposes gpiochip0-4; verify with `gpiodetect`
BUTTON_LINE_OFFSET = 17               # Change to match the physical pin you wired the button to

AIRTABLE_TOKEN = os.getenv("AIRTABLE_TOKEN")
AIRTABLE_BASE_ID = os.getenv("AIRTABLE_BASE_ID")
AIRTABLE_TABLE_NAME = os.getenv("AIRTABLE_TABLE_NAME", "Patients")
AIRTABLE_ID_FIELD = os.getenv("AIRTABLE_ID_FIELD", "ID")  # name of the column holding the patient ID

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

VOSK_MODEL_PATH = os.getenv(
    "VOSK_MODEL_PATH",
    os.path.join(os.path.dirname(os.path.abspath(__file__)), "models", "vosk-model-small-en-us-0.15"),
)
SetLogLevel(-1)  # silence Vosk's default debug logging

REPORT_UPLOAD_URL = "https://nexus-medi-backend-gykw.onrender.com/api/v1/reports"
REPORT_UPLOAD_AUTH = "nexusrobogenn0825"

# ============================================================
# TRIGGER SYSTEM: supports BOTH a physical button AND keyboard Enter
# ============================================================
def setup_button():
    """
    Sets up the button using libgpiod instead of RPi.GPIO.
    Returns a gpiod line request object, or None if no button hardware
    is present/configured (in which case only the keyboard trigger works).
    """
    try:
        chip = gpiod.Chip(BUTTON_GPIO_CHIP)
        line = chip.get_line(BUTTON_LINE_OFFSET)
        line.request(consumer="nexus-robot", type=gpiod.LINE_REQ_DIR_IN, flags=gpiod.LINE_REQ_FLAG_BIAS_PULL_UP)
        print("Physical button detected and ready.")
        return line
    except Exception as e:
        print(f"No physical button available ({e}). Keyboard (Enter) trigger only.")
        return None


def button_watcher(line, trigger_queue, poll_interval=0.05):
    """Runs forever in a background thread. Puts ('button','press'/'release') on the queue."""
    if line is None:
        return
    prev_state = 1  # released, since pull-up means idle HIGH
    while True:
        try:
            value = line.get_value()
        except Exception:
            return
        if value != prev_state:
            trigger_queue.put(('button', 'press' if value == 0 else 'release'))
            prev_state = value
            time.sleep(0.2)  # simple debounce
        time.sleep(poll_interval)


def keyboard_watcher(trigger_queue):
    """Runs forever in a background thread. Puts ('keyboard','press') on the queue on every Enter."""
    while True:
        try:
            input()
            trigger_queue.put(('keyboard', 'press'))
        except EOFError:
            return


def wait_for_start_trigger(trigger_queue):
    """Blocks until either the button is pressed or Enter is hit. Returns the source: 'button' or 'keyboard'."""
    while True:
        source, action = trigger_queue.get()
        if action == 'press':
            return source


def wait_for_stop_trigger(trigger_queue, source):
    """
    Blocks until the session is stopped using the SAME method that started it:
    - if started by button, waits for the button release
    - if started by keyboard, waits for another Enter press
    """
    while True:
        s, a = trigger_queue.get()
        if source == 'button' and s == 'button' and a == 'release':
            return
        if source == 'keyboard' and s == 'keyboard' and a == 'press':
            return


# ============================================================
# AIRTABLE: fetch patient ID from first row
# ============================================================
def get_patient_id_from_airtable():
    """
    Fetches the patient ID from the first row of the configured Airtable table.
    """
    if not (AIRTABLE_TOKEN and AIRTABLE_BASE_ID):
        print("Airtable credentials missing in .env")
        return None

    url = f"https://api.airtable.com/v0/{AIRTABLE_BASE_ID}/{AIRTABLE_TABLE_NAME}"
    headers = {"Authorization": f"Bearer {AIRTABLE_TOKEN}"}
    params = {"maxRecords": 1, "view": "Grid view"}  # adjust view name if yours differs

    try:
        response = requests.get(url, headers=headers, params=params, timeout=10)
        response.raise_for_status()
        data = response.json()
        records = data.get("records", [])
        if not records:
            print("No records found in Airtable table.")
            return None

        fields = records[0].get("fields", {})
        patient_id = fields.get(AIRTABLE_ID_FIELD)
        if not patient_id:
            print(f"Field '{AIRTABLE_ID_FIELD}' not found in first record.")
            return None

        return str(patient_id)

    except requests.exceptions.RequestException as e:
        print(f"Error fetching patient ID from Airtable: {e}")
        return None


# ============================================================
# MICROPHONE: auto-detect USB mic
# ============================================================
def find_usb_microphone_index():
    """
    Scans available microphones and returns the device index of the first
    one whose name contains 'USB' (case-insensitive). Falls back to the
    default microphone if none found.
    """
    mic_names = sr.Microphone.list_microphone_names()
    for index, name in enumerate(mic_names):
        if "usb" in name.lower():
            print(f"Using USB microphone: [{index}] {name}")
            return index

    print("No USB microphone found by name — falling back to default input device.")
    return None  # sr.Microphone() with device_index=None uses system default


# ============================================================
# AUDIO / SPEECH
# ============================================================
def speech_to_text_continuous(stop_event, mic_index, vosk_model):
    audio_queue = Queue()
    recognizer = sr.Recognizer()
    combined_text = []

    def record_audio():
        with sr.Microphone(device_index=mic_index) as source:
            recognizer.adjust_for_ambient_noise(source)
            print("Recording... (release button to stop)")

            while not stop_event.is_set():
                try:
                    audio = recognizer.listen(source, timeout=0.1, phrase_time_limit=10)
                    audio_queue.put(audio)
                except sr.WaitTimeoutError:
                    continue
                except Exception as e:
                    print(f"Recording error: {e}")
                    break

    def process_audio():
        while not stop_event.is_set() or not audio_queue.empty():
            try:
                audio = audio_queue.get(timeout=1)
                raw_data = audio.get_raw_data(convert_rate=16000, convert_width=2)
                chunk_recognizer = KaldiRecognizer(vosk_model, 16000)
                chunk_recognizer.AcceptWaveform(raw_data)
                result = json.loads(chunk_recognizer.FinalResult())
                text = result.get("text", "").strip()
                if text:
                    combined_text.append(text)
                    print(f"Partial: {text}")
            except Empty:
                continue

    record_thread = threading.Thread(target=record_audio, daemon=True)
    process_thread = threading.Thread(target=process_audio, daemon=True)

    record_thread.start()
    process_thread.start()

    while not stop_event.is_set():
        time.sleep(0.1)

    print("\nStopping recording...")
    record_thread.join(timeout=2.0)
    process_thread.join(timeout=2.0)

    return " ".join(combined_text)


def read_text(text):
    """Text-to-speech playback via the espeak CLI (pyttsx3's espeak driver is
    incompatible with current espeak-ng voice-list ABI on recent Debian)."""
    print(text)
    try:
        subprocess.run(["espeak", "-s", "150", text], check=True)
    except (subprocess.CalledProcessError, FileNotFoundError) as e:
        print(f"TTS playback failed: {e}")


# ============================================================
# GEMINI REPORT GENERATION
# ============================================================
def generate(user_input, chat_history, medical_prompt):
    client = genai.Client(api_key=GEMINI_API_KEY)
    model = "gemini-2.5-flash"

    contents = [types.Content(role="user", parts=[types.Part.from_text(text=medical_prompt)])]

    for message in chat_history:
        contents.append(types.Content(role=message["role"], parts=[types.Part.from_text(text=message["text"])]))

    contents.append(types.Content(role="user", parts=[types.Part.from_text(text=user_input)]))

    tools = [types.Tool(google_search=types.GoogleSearch())]
    generate_content_config = types.GenerateContentConfig(
        temperature=1,
        top_p=0.95,
        top_k=40,
        max_output_tokens=8192,
        tools=tools,
        response_mime_type="text/plain",
    )

    max_attempts = 3
    retry_delay_seconds = 10
    for attempt in range(1, max_attempts + 1):
        try:
            response_text = ""
            for chunk in client.models.generate_content_stream(
                model=model,
                contents=contents,
                config=generate_content_config,
            ):
                print(chunk.text, end="")
                response_text += chunk.text
            return response_text
        except genai_errors.ServerError as e:
            print(f"\nGemini server error (attempt {attempt}/{max_attempts}): {e}")
            if attempt == max_attempts:
                raise
            time.sleep(retry_delay_seconds)


# ============================================================
# REPORT UPLOAD / SAVE
# ============================================================
def ping_backend():
    """
    Fires a lightweight request at the backend as soon as a session starts.
    The backend (Render free tier) sleeps when idle and takes a while to wake
    up on the first request, so pinging it early means it's already warm by
    the time the session ends and the real upload happens.
    """
    try:
        base_url = urlparse(REPORT_UPLOAD_URL)
        requests.get(f"{base_url.scheme}://{base_url.netloc}/", timeout=15)
        print("Backend ping sent.")
    except requests.exceptions.RequestException as e:
        print(f"Backend ping failed (will still attempt upload later): {e}")


def upload_report_to_server(report_text, patient_id):
    if not patient_id:
        print("No patient ID available for upload")
        return False

    headers = {
        'Content-Type': 'application/json',
        'Authorization': REPORT_UPLOAD_AUTH
    }
    payload = {
        "patient_id": patient_id,
        "report_summary": report_text,
        "created_by": "Robot Nexus",
        "isconfidential": True,
        "status": "Not Responded"
    }

    try:
        print("\nUploading report to server...")
        response = requests.post(REPORT_UPLOAD_URL, headers=headers, data=json.dumps(payload), timeout=10)
        response.raise_for_status()
        result = response.json()
        print("\n✅ Report uploaded successfully!")
        print(f"Report ID: {result['data']['report_id']}")
        print(f"Status: {result['data']['status']}")
        return True

    except requests.exceptions.RequestException as e:
        print("\n❌ Error uploading report:")
        if hasattr(e, 'response') and e.response is not None:
            try:
                error_data = e.response.json()
                print(f"Status Code: {e.response.status_code}")
                print(f"Error: {error_data.get('error', 'Unknown error')}")
            except ValueError:
                print(f"Status Code: {e.response.status_code}")
                print(f"Response Text: {e.response.text}")
        else:
            print(f"Connection Error: {str(e)}")
        return False


def save_report_to_word(report_text):
    doc = Document()
    doc.add_heading("Consultation Report", 0)

    for line in report_text.split('\n'):
        if line.strip():
            doc.add_paragraph(line.strip())

    script_dir = os.path.dirname(os.path.abspath(__file__))
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    filename = f"Consultation_Report_{timestamp}.docx"
    full_path = os.path.join(script_dir, filename)

    doc.save(full_path)
    print(f"\nReport saved as: {full_path}")


def save_transcript_fallback(transcript_text):
    """Saves the raw transcript as plain text when report generation fails, so the consultation isn't lost."""
    script_dir = os.path.dirname(os.path.abspath(__file__))
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    filename = f"Transcript_{timestamp}.txt"
    full_path = os.path.join(script_dir, filename)

    with open(full_path, "w", encoding="utf-8") as f:
        f.write(transcript_text)

    print(f"\nRaw transcript saved as: {full_path}")
    return full_path


# ============================================================
# MAIN LOOP
# ============================================================
def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    prompt_path = os.path.join(script_dir, "medical_prompt.txt")
    with open(prompt_path, "r", encoding="utf-8") as file:
        medical_prompt = file.read()

    chat_history = []
    mic_index = find_usb_microphone_index()
    button_line = setup_button()

    print(f"Loading speech recognition model from {VOSK_MODEL_PATH} ...")
    vosk_model = Model(VOSK_MODEL_PATH)
    print("Speech recognition model loaded.")

    trigger_queue = Queue()
    threading.Thread(target=button_watcher, args=(button_line, trigger_queue), daemon=True).start()
    threading.Thread(target=keyboard_watcher, args=(trigger_queue,), daemon=True).start()

    try:
        read_text("System ready. Press the button or hit Enter to start a consultation session.")

        while True:
            trigger_source = wait_for_start_trigger(trigger_queue)

            threading.Thread(target=ping_backend, daemon=True).start()

            patient_id = get_patient_id_from_airtable()
            if not patient_id:
                read_text("Failed to fetch patient ID from Airtable. Please check the database.")
                if trigger_source == 'button':
                    wait_for_stop_trigger(trigger_queue, 'button')  # clear the pending release
                continue

            stop_event = threading.Event()
            read_text("Recording in progress.")

            result_container = {}

            def record_and_store():
                result_container['text'] = speech_to_text_continuous(stop_event, mic_index, vosk_model)

            recording_thread = threading.Thread(target=record_and_store, daemon=True)
            recording_thread.start()

            wait_for_stop_trigger(trigger_queue, trigger_source)

            stop_event.set()
            recording_thread.join()

            read_text("Session ended. Generating report...")

            consultation_text = result_container.get('text', '')
            if consultation_text:
                try:
                    response_text = generate(consultation_text, chat_history, medical_prompt)
                except Exception as e:
                    print(f"\nReport generation failed: {e}")
                    save_transcript_fallback(consultation_text)
                    read_text("Report generation failed. The raw transcript has been saved locally.")
                    continue

                save_report_to_word(response_text)

                if upload_report_to_server(response_text, patient_id):
                    read_text("Report has been successfully uploaded to the patient's records.")
                else:
                    read_text("Warning: Could not upload report to server. Local copy has been saved.")

                chat_history.append({"role": "user", "text": consultation_text})
                chat_history.append({"role": "model", "text": response_text})
            else:
                read_text("No speech was captured during the session.")

    except KeyboardInterrupt:
        read_text("Exiting...")
    finally:
        if button_line is not None:
            button_line.release()


if __name__ == "__main__":
    main()
